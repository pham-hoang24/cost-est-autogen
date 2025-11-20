#!/usr/bin/env python3
"""
Data Anonymization Service - Clean Working Version with Real AI Integration
"""

from fastapi import FastAPI, UploadFile, File, HTTPException, Request, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.security import HTTPBearer
from fastapi.middleware.trustedhost import TrustedHostMiddleware
import uvicorn
import os
import re
import uuid
import io
import csv
import json
from datetime import datetime
from typing import Dict, List, Any, Union
import PyPDF2
from docx import Document
import openpyxl
from faker import Faker
import hashlib
import tempfile
import nltk
import usaddress
from dateutil import parser as date_parser
import aiohttp
import asyncio
from .gdpr_anonymizer import get_gdpr_anonymizer

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

# Configuration constants
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB
ALLOWED_EXTENSIONS = {'.csv', '.xlsx', '.xls', '.pdf', '.docx', '.txt', '.json'}
RATE_LIMIT_REQUESTS = 100  # requests per minute
RATE_LIMIT_WINDOW = 60  # seconds

# Initialize FastAPI app
app = FastAPI(
    title="GDPR-Compliant Data Anonymization Service",
    description="Professional data anonymization microservice with Microsoft Presidio integration",
    version="2.0.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

# Security
security = HTTPBearer(auto_error=False)

# CORS configuration
cors_origins = os.getenv("CORS_ORIGINS", "http://localhost:3000,http://localhost:3001").split(",")
print(f"🔧 CORS Origins: {cors_origins}")
print(f"🔧 Environment CORS_ORIGINS: {os.getenv('CORS_ORIGINS')}")

# Add explicit CORS origins for debugging - including all possible frontend URLs
cors_origins = [
    "http://localhost:3000",   # SW4E UI dev
    "http://localhost:3001",   # SW4E UI production 
    "http://localhost:3002",   # Data Anonymization Frontend
    "http://127.0.0.1:3000",
    "http://127.0.0.1:3001", 
    "http://127.0.0.1:3002",
    "*"  # Allow all origins for development
]
print(f"🔧 Final CORS Origins: {cors_origins}")

# Add security middleware
app.add_middleware(
    TrustedHostMiddleware, 
    allowed_hosts=["*"]
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=cors_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition", "Content-Type", "Content-Length"]
)

# Rate limiting storage
rate_limit_storage = {}

def check_rate_limit(request: Request):
    """Simple rate limiting implementation"""
    client_ip = request.client.host
    current_time = datetime.now().timestamp()
    
    if client_ip not in rate_limit_storage:
        rate_limit_storage[client_ip] = []
    
    # Clean old requests
    rate_limit_storage[client_ip] = [
        req_time for req_time in rate_limit_storage[client_ip] 
        if current_time - req_time < RATE_LIMIT_WINDOW
    ]
    
    # Check if limit exceeded
    if len(rate_limit_storage[client_ip]) >= RATE_LIMIT_REQUESTS:
        raise HTTPException(
            status_code=429, 
            detail="Rate limit exceeded. Please try again later."
        )
    
    # Add current request
    rate_limit_storage[client_ip].append(current_time)

def validate_file_size(file: UploadFile):
    """Validate file size"""
    if hasattr(file, 'size') and file.size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"File too large. Maximum size allowed is {MAX_FILE_SIZE // (1024*1024)}MB"
        )

def validate_file_extension(filename: str):
    """Validate file extension"""
    file_ext = os.path.splitext(filename)[1].lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type not supported. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
        )
    return file_ext

# Initialize Faker
fake = Faker()

# In-memory storage
uploaded_files = {}
processing_status = {}  # Track processing progress

# AI Configuration with OpenRouter support
ai_config = {
    "enabled": True,
    "provider": "openrouter",  # openrouter, openai, anthropic, etc.
    "model": "anthropic/claude-3-haiku",  # Default model
    "api_key": None,
    "base_url": "https://openrouter.ai/api/v1",  # OpenRouter API endpoint
    "temperature": 0.3,
    "max_tokens": 150,
    "fallback_to_traditional": True  # Always fallback to traditional if AI fails
}

class ContextAnalyzer:
    """Analyzes context around sensitive data for better anonymization"""
    
    def __init__(self):
        self.context_patterns = {
            'name': {
                'titles': r'\b(?:Mr|Mrs|Ms|Dr|Prof|Sir|Madam|President|CEO|Manager|Director)\b',
                'roles': r'\b(?:employee|manager|director|executive|staff|member|user|customer|client|patient)\b',
                'organizations': r'\b(?:company|corporation|inc|llc|ltd|corp|organization|institution|hospital|clinic|school|university)\b'
            },
            'address': {
                'location_types': r'\b(?:residence|home|office|workplace|facility|building|complex|center|plaza|mall)\b',
                'geographic': r'\b(?:city|town|village|county|state|province|country|region|area|district)\b'
            },
            'financial': {
                'amounts': r'\b(?:salary|wage|income|revenue|profit|cost|price|fee|charge|payment|transaction)\b',
                'currencies': r'\b(?:dollar|euro|pound|yen|rupee|peso|franc|mark|lira|dinar)\b'
            },
            'medical': {
                'conditions': r'\b(?:diagnosis|condition|symptom|treatment|medication|prescription|test|result|report|record)\b',
                'professions': r'\b(?:doctor|physician|nurse|specialist|therapist|surgeon|dentist|pharmacist)\b'
            }
        }
    
    def analyze_csv_context(self, headers: List[str], row_data: List[str]) -> dict:
        """Analyze CSV context for better anonymization"""
        context_info = {
            'domain': 'corporate',  # Default for employee data
            'formality': 'formal',
            'urgency': 'normal',
            'relationships': 'employee',
            'industry': 'corporate',
            'column_context': {}
        }
        
        # Analyze column headers for context
        for i, header in enumerate(headers):
            header_lower = header.lower()
            if 'salary' in header_lower or 'commission' in header_lower:
                context_info['column_context'][i] = 'financial'
            elif 'hire_date' in header_lower or 'date' in header_lower:
                context_info['column_context'][i] = 'temporal'
            elif 'email' in header_lower:
                context_info['column_context'][i] = 'communication'
            elif 'phone' in header_lower:
                context_info['column_context'][i] = 'communication'
            elif 'name' in header_lower:
                context_info['column_context'][i] = 'personal'
            elif 'id' in header_lower:
                context_info['column_context'][i] = 'identifier'
            elif 'job' in header_lower or 'department' in header_lower:
                context_info['column_context'][i] = 'organizational'
        
        # Detect domain from data patterns
        if any('salary' in str(cell).lower() for cell in row_data):
            context_info['domain'] = 'corporate'
        elif any('patient' in str(cell).lower() for cell in row_data):
            context_info['domain'] = 'medical'
        elif any('account' in str(cell).lower() for cell in row_data):
            context_info['domain'] = 'financial'
        
        return context_info
    
    def analyze_context(self, text: str, surrounding_context: str = "") -> dict:
        """Analyze context around sensitive data"""
        context_info = {
            'domain': self._detect_domain(text, surrounding_context),
            'formality': self._detect_formality(text),
            'urgency': self._detect_urgency(surrounding_context),
            'relationships': self._detect_relationships(surrounding_context),
            'industry': self._detect_industry(surrounding_context)
        }
        return context_info
    
    def _detect_domain(self, text: str, context: str) -> str:
        """Detect the domain/industry context"""
        full_text = f"{text} {context}".lower()
        
        if any(word in full_text for word in ['medical', 'health', 'patient', 'diagnosis', 'treatment']):
            return 'medical'
        elif any(word in full_text for word in ['financial', 'bank', 'account', 'transaction', 'payment']):
            return 'financial'
        elif any(word in full_text for word in ['legal', 'case', 'client', 'attorney', 'court']):
            return 'legal'
        elif any(word in full_text for word in ['education', 'student', 'teacher', 'school', 'university']):
            return 'education'
        elif any(word in full_text for word in ['corporate', 'business', 'company', 'employee', 'manager']):
            return 'corporate'
        else:
            return 'general'
    
    def _detect_formality(self, text: str) -> str:
        """Detect formality level of the text"""
        formal_indicators = ['respectfully', 'sincerely', 'regards', 'dear', 'to whom it may concern']
        informal_indicators = ['hey', 'hi', 'thanks', 'cool', 'awesome']
        
        text_lower = text.lower()
        formal_count = sum(1 for indicator in formal_indicators if indicator in text_lower)
        informal_count = sum(1 for indicator in informal_indicators if indicator in text_lower)
        
        if formal_count > informal_count:
            return 'formal'
        elif informal_count > formal_count:
            return 'informal'
        else:
            return 'neutral'
    
    def _detect_urgency(self, context: str) -> str:
        """Detect urgency level in context"""
        urgent_words = ['urgent', 'emergency', 'immediate', 'asap', 'critical', 'urgent']
        context_lower = context.lower()
        
        if any(word in context_lower for word in urgent_words):
            return 'high'
        else:
            return 'normal'
    
    def _detect_relationships(self, context: str) -> str:
        """Detect relationships mentioned in context"""
        relationship_words = ['family', 'parent', 'child', 'spouse', 'partner', 'colleague', 'friend']
        context_lower = context.lower()
        
        for word in relationship_words:
            if word in context_lower:
                return word
        return 'unknown'
    
    def _detect_industry(self, context: str) -> str:
        """Detect industry from context"""
        industry_keywords = {
            'healthcare': ['medical', 'health', 'hospital', 'clinic', 'patient'],
            'finance': ['bank', 'financial', 'investment', 'insurance', 'accounting'],
            'technology': ['software', 'tech', 'digital', 'computer', 'programming'],
            'education': ['school', 'university', 'college', 'learning', 'academic'],
            'legal': ['law', 'legal', 'attorney', 'court', 'case'],
            'retail': ['store', 'shop', 'retail', 'commerce', 'customer']
        }
        
        context_lower = context.lower()
        for industry, keywords in industry_keywords.items():
            if any(keyword in context_lower for keyword in keywords):
                return industry
        return 'general'

class SemanticMapper:
    """Generates semantically appropriate replacements based on context"""
    
    def __init__(self):
        self.fake = Faker()
        self.domain_templates = {
            'medical': {
                'name': ['Dr. {first} {last}', 'Prof. {first} {last}', 'Dr. {last}'],
                'address': ['Medical Center {number}', 'Healthcare Plaza {number}', 'Medical District {number}'],
                'email': ['dr.{last}@medical.org', 'prof.{last}@healthcare.edu', 'physician.{last}@clinic.com']
            },
            'financial': {
                'name': ['Mr. {first} {last}', 'Ms. {first} {last}', 'Director {last}'],
                'address': ['Financial District {number}', 'Business Plaza {number}', 'Corporate Center {number}'],
                'email': ['{first}.{last}@finance.com', 'director.{last}@corp.com', 'executive.{last}@bank.com']
            },
            'corporate': {
                'name': ['{first} {last}', 'Mr. {last}', 'Ms. {last}'],
                'address': ['Office {number}', 'Business Park {number}', 'Corporate Plaza {number}'],
                'email': ['{first}.{last}@company.com', '{last}@corp.com', 'employee.{last}@business.com']
            },
            'general': {
                'name': ['{first} {last}', 'Person {number}', 'Individual {number}'],
                'address': ['Location {number}', 'Address {number}', 'Place {number}'],
                'email': ['user{number}@example.com', 'person{number}@domain.com', 'individual{number}@mail.com']
            }
        }
    
    def generate_replacement(self, field_type: str, original_text: str, context_info: dict) -> str:
        """Generate contextually appropriate replacement"""
        domain = context_info.get('domain', 'general')
        formality = context_info.get('formality', 'neutral')
        
        if field_type == 'name':
            return self._generate_name_replacement(domain, formality)
        elif field_type == 'address':
            return self._generate_address_replacement(domain, formality)
        elif field_type == 'email':
            return self._generate_email_replacement(domain, formality)
        elif field_type == 'phone':
            return self._generate_phone_replacement(domain)
        elif field_type == 'date':
            return self._generate_date_replacement(context_info)
        elif field_type == 'postal_code':
            return self._generate_postal_code_replacement(domain)
        elif field_type == 'salary':
            return self._generate_salary_replacement(domain)
        elif field_type == 'id':
            return self._generate_id_replacement(domain)
        else:
            return self._generate_generic_replacement(field_type, domain)
    
    def _generate_name_replacement(self, domain: str, formality: str) -> str:
        """Generate contextually appropriate name replacement"""
        templates = self.domain_templates.get(domain, self.domain_templates['general'])
        name_template = templates['name'][0] if templates['name'] else '{first} {last}'
        
        if formality == 'formal':
            # Use more formal titles
            if 'Dr.' in name_template or 'Prof.' in name_template:
                return name_template.format(
                    first=self.fake.first_name(),
                    last=self.fake.last_name()
                )
            else:
                return f"Mr. {self.fake.first_name()} {self.fake.last_name()}"
        else:
            return name_template.format(
                first=self.fake.first_name(),
                last=self.fake.last_name()
            )
    
    def _generate_address_replacement(self, domain: str, formality: str) -> str:
        """Generate contextually appropriate address replacement"""
        templates = self.domain_templates.get(domain, self.domain_templates['general'])
        address_template = templates['address'][0] if templates['address'] else 'Location {number}'
        
        street_number = self.fake.building_number()
        street_name = self.fake.street_name()
        city = self.fake.city()
        state = self.fake.state_abbr()
        zip_code = self.fake.zipcode()
        
        if domain == 'medical':
            return f"Medical Center {street_number}, {city}, {state} {zip_code}"
        elif domain == 'financial':
            return f"Financial District {street_number}, {city}, {state} {zip_code}"
        else:
            return f"{street_number} {street_name}, {city}, {state} {zip_code}"
    
    def _generate_email_replacement(self, domain: str, formality: str) -> str:
        """Generate contextually appropriate email replacement"""
        templates = self.domain_templates.get(domain, self.domain_templates['general'])
        email_template = templates['email'][0] if templates['email'] else 'user{number}@example.com'
        
        if domain == 'medical':
            return f"dr.{self.fake.last_name()}@medical.org"
        elif domain == 'financial':
            return f"{self.fake.first_name()}.{self.fake.last_name()}@finance.com"
        else:
            return f"user{self.fake.random_number(digits=4)}@example.com"
    
    def _generate_phone_replacement(self, domain: str) -> str:
        """Generate contextually appropriate phone replacement"""
        if domain == 'medical':
            # Medical offices often have specific area codes
            return f"555-{self.fake.random_number(digits=3)}-{self.fake.random_number(digits=4)}"
        else:
            return f"{self.fake.random_number(digits=3)}-{self.fake.random_number(digits=3)}-{self.fake.random_number(digits=4)}"
    
    def _generate_date_replacement(self, context_info: dict) -> str:
        """Generate contextually appropriate date replacement"""
        import random
        from datetime import datetime, timedelta
        
        base_date = datetime.now()
        if context_info.get('urgency') == 'high':
            # For urgent contexts, use recent dates
            days_offset = random.randint(-30, 0)
        else:
            # For normal contexts, use varied dates
            days_offset = random.randint(-365, 365)
        
        new_date = base_date + timedelta(days=days_offset)
        return new_date.strftime("%m/%d/%Y")
    
    def _generate_generic_replacement(self, field_type: str, domain: str) -> str:
        """Generate generic replacement for unknown field types"""
        if field_type == 'ssn':
            return f"***-**-{self.fake.random_number(digits=4)}"
        elif field_type == 'credit_card':
            return f"**** **** **** {self.fake.random_number(digits=4)}"
        elif field_type == 'ip_address':
            return f"{self.fake.random_number(digits=1, fix_len=True)}.{self.fake.random_number(digits=1, fix_len=True)}.*.*"
        else:
            return f"[ANONYMIZED_{field_type.upper()}]"
    
    def _generate_postal_code_replacement(self, domain: str) -> str:
        """Generate contextually appropriate postal code replacement"""
        if domain == 'medical':
            return f"MC{self.fake.random_number(digits=5)}"
        elif domain == 'financial':
            return f"FC{self.fake.random_number(digits=5)}"
        elif domain == 'corporate':
            return f"CC{self.fake.random_number(digits=5)}"
        else:
            return f"PC{self.fake.random_number(digits=5)}"
    
    def _generate_salary_replacement(self, domain: str) -> str:
        """Generate contextually appropriate salary replacement"""
        if domain == 'medical':
            base_salary = self.fake.random_int(min=80000, max=250000)
        elif domain == 'financial':
            base_salary = self.fake.random_int(min=60000, max=200000)
        elif domain == 'corporate':
            base_salary = self.fake.random_int(min=40000, max=150000)
        else:
            base_salary = self.fake.random_int(min=30000, max=100000)
        
        return str(base_salary)
    
    def _generate_id_replacement(self, domain: str) -> str:
        """Generate contextually appropriate ID replacement"""
        if domain == 'medical':
            return f"MED{self.fake.random_number(digits=6)}"
        elif domain == 'financial':
            return f"FIN{self.fake.random_number(digits=6)}"
        elif domain == 'corporate':
            return f"EMP{self.fake.random_number(digits=6)}"
        else:
            return f"ID{self.fake.random_number(digits=6)}"

# Clean regex patterns with NO capture groups - ENHANCED VERSIONS
SENSITIVE_PATTERNS = {
    'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
    'phone': r'\b(?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}\b',
    'phone_alt': r'\b[0-9]{3}[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}\b',  # Alternative format
    'ssn': r'\b\d{3}[-\s]?\d{2}[-\s]?\d{4}\b',
    'eu_ssn': r'\b(?:\d{9}|\d{10}|\d{11}|\d{12}|\d{13}|\d{15}|[A-Z]{2}\d{6}[A-Z]|[A-Z]{6}\d{2}[A-Z]\d{2}[A-Z]\d{3}[A-Z]|\d{6}[-]?\d{4}|\d{6}[A-Z]\d{3}[A-Z0-9]|\d{6}[/]?\d{4}|\d{7}[A-Z]{1,2})\b',
    'credit_card': r'\b(?:4[0-9]{12}(?:[0-9]{3})?|5[1-5][0-9]{14}|3[47][0-9]{13}|3[0-9]{13}|6(?:011|5[0-9]{2})[0-9]{12})\b',
    'credit_card_exp': r'\b(?:0[1-9]|1[0-2])[/-]\d{2}\b',
    'passport': r'\b[A-Z]{2,3}\d{6,9}\b',
    'name': r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b',
    'address': r'\b\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr|Way|Court|Ct|Place|Pl|Square|Sq|Circle|Cir|Terrace|Ter)\b',
    'postal_code': r'\b\d{5}(?:-\d{4})?\b',
    'ip_address': r'\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\b',
    'bank_account': r'\b\d{8,17}\b',
    'date': r'\b(?:0?[1-9]|1[0-2])[/-](?:0?[1-9]|[12]\d|3[01])[/-](?:19|20)\d{2}\b',
    'mac_address': r'\b(?:[0-9A-Fa-f]{2}[:-]){5}[0-9A-Fa-f]{2}\b',
    'vehicle_plate': r'\b[A-Z]{1,3}[0-9]{1,4}[A-Z]{1,2}\b',
    # CV-specific patterns
    'cv_name': r'\b(?:Name|Full Name|First Name|Last Name|Given Name|Surname)[\s:]*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
    'cv_email': r'\b(?:Email|E-mail|Contact)[\s:]*([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})',
    'cv_phone': r'\b(?:Phone|Tel|Mobile|Cell|Contact)[\s:]*((?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4})',
    'cv_address': r'\b(?:Address|Location|Residence)[\s:]*(\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr|Way|Court|Ct|Place|Pl|Square|Sq|Circle|Cir|Terrace|Ter))',
    'cv_linkedin': r'\b(?:LinkedIn|LinkedIn Profile)[\s:]*([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}|linkedin\.com/in/[A-Za-z0-9_-]+)',
    'cv_github': r'\b(?:GitHub|Github|GitHub Profile)[\s:]*(github\.com/[A-Za-z0-9_-]+)',
    'cv_website': r'\b(?:Website|Portfolio|Personal Website)[\s:]*(https?://[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})',
    'cv_skills': r'\b(?:Skills|Technical Skills|Core Competencies)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_experience': r'\b(?:Experience|Work Experience|Professional Experience)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_education': r'\b(?:Education|Academic Background|Qualifications)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_certification': r'\b(?:Certifications|Certificates|Licenses)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_project': r'\b(?:Projects|Portfolio|Key Projects)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_achievement': r'\b(?:Achievements|Awards|Honors)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_reference': r'\b(?:References|Referees)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_objective': r'\b(?:Objective|Summary|Profile|About)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_interest': r'\b(?:Interests|Hobbies|Personal Interests)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_language': r'\b(?:Languages|Language Skills)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_availability': r'\b(?:Availability|Notice Period|Start Date)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_salary': r'\b(?:Salary|Compensation|Expected Salary|Salary Expectation)[\s:]*(\$?[0-9,]+(?:\.\d{2})?)',
    'cv_location': r'\b(?:Location|Current Location|Based in)[\s:]*([A-Za-z\s,]+)',
    'cv_nationality': r'\b(?:Nationality|Citizenship|Visa Status)[\s:]*([A-Za-z\s,]+)',
    'cv_dob': r'\b(?:Date of Birth|DOB|Born)[\s:]*((?:0?[1-9]|1[0-2])[/-](?:0?[1-9]|[12]\d|3[01])[/-](?:19|20)\d{2})',
    'cv_gender': r'\b(?:Gender|Sex)[\s:]*([A-Za-z]+)',
    'cv_marital': r'\b(?:Marital Status|Relationship Status)[\s:]*([A-Za-z\s]+)',
    'cv_driving': r'\b(?:Driving License|Driver\'s License|License Number)[\s:]*([A-Za-z0-9\s-]+)',
    'cv_passport': r'\b(?:Passport|Passport Number)[\s:]*([A-Z]{2,3}\d{6,9})',
    'cv_id': r'\b(?:ID|ID Number|Identification)[\s:]*([A-Za-z0-9\s-]+)',
    'cv_ssn': r'\b(?:SSN|Social Security|Social Security Number)[\s:]*(\d{3}[-\s]?\d{2}[-\s]?\d{4})',
    'cv_credit': r'\b(?:Credit Score|Credit Rating)[\s:]*([0-9]+)',
    'cv_bank': r'\b(?:Bank Account|Account Number)[\s:]*(\d{8,17})',
    'cv_insurance': r'\b(?:Insurance|Insurance Number|Policy Number)[\s:]*([A-Za-z0-9\s-]+)',
    'cv_medical': r'\b(?:Medical|Health|Medical History)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_criminal': r'\b(?:Criminal|Criminal Record|Background Check)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_emergency': r'\b(?:Emergency Contact|Emergency)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_family': r'\b(?:Family|Dependents|Children)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_religion': r'\b(?:Religion|Religious Affiliation)[\s:]*([A-Za-z\s]+)',
    'cv_political': r'\b(?:Political|Political Affiliation)[\s:]*([A-Za-z\s]+)',
    'employee_id': r'\bEMP-\d{4}-\d{3}\b',
    'salary': r'\$\d{1,3}(?:,\d{3})*(?:\.\d{2})?',
    'department': r'\b(?:Human Resources|Engineering|Marketing|Sales|Finance|Operations|IT|Legal|HR|Admin|Administration)\b',
    'routing_number': r'\b(?:021000021|111000025|123456789)\b',
    'cv_disability': r'\b(?:Disability|Special Needs)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_veteran': r'\b(?:Veteran|Military Service)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_union': r'\b(?:Union|Union Membership)[\s:]*([A-Za-z0-9\s,.-]+)',
    'cv_other': r'\b(?:Other|Additional|Miscellaneous)[\s:]*([A-Za-z0-9\s,.-]+)'
}

class AnonymizationEngine:
    def __init__(self):
        self.fake = Faker()
        self.entity_map = {}
        self.counter = 0
        self.context_analyzer = ContextAnalyzer()
        self.semantic_mapper = SemanticMapper()
    
    def ai_anonymize_with_context(self, text: str, field_type: str, context: str = "") -> str:
        """AI-powered anonymization with context awareness"""
        try:
            # Analyze context for better anonymization
            context_info = self.context_analyzer.analyze_context(text, context)
            
            # Generate semantically appropriate replacement
            replacement = self.semantic_mapper.generate_replacement(
                field_type, text, context_info
            )
            
            # Ensure consistency across the document
            key = f"{field_type}:{text.lower()}"
            if key not in self.entity_map:
                self.entity_map[key] = replacement
            else:
                replacement = self.entity_map[key]  # Use existing mapping
            
            return replacement
            
        except Exception as e:
            print(f"AI anonymization failed: {e}")
            # Fallback to traditional method
            return self._traditional_anonymize(field_type, text)
    
    def _traditional_anonymize(self, field_type: str, text: str) -> str:
        """Traditional anonymization as fallback"""
        if field_type == "name":
            return self.anonymize_name(text)
        elif field_type == "email":
            return self.anonymize_email(text)
        elif field_type == "phone":
            return self.anonymize_phone(text)
        elif field_type == "ssn" or field_type == "cv_ssn":
            return self.anonymize_ssn(text)
        elif field_type == "credit_card":
            return self.anonymize_credit_card(text)
        elif field_type == "address":
            return self.anonymize_address(text)
        elif field_type == "date":
            return self.anonymize_date(text)
        elif field_type == "ip_address":
            return self.anonymize_ip_address(text)
        elif field_type == "salary" or field_type == "cv_salary":
            return self.anonymize_salary(text)
        elif field_type == "employee_id" or field_type == "cv_id":
            return self.anonymize_employee_id(text)
        elif field_type == "department":
            return self.anonymize_department(text)
        elif field_type == "routing_number":
            return self.anonymize_routing_number(text)
        elif field_type == "bank_account":
            return self.anonymize_bank_account(text)
        elif field_type == "passport":
            return self.anonymize_passport(text)
        elif field_type == "eu_ssn":
            return self.anonymize_eu_ssn(text)
        else:
            return f"[ANONYMIZED_{field_type.upper()}]"
    
    def get_consistent_pseudonym(self, entity_type: str, value: str) -> str:
        key = f"{entity_type}:{value.lower()}"
        if key not in self.entity_map:
            self.counter += 1
            if entity_type == "person":
                self.entity_map[key] = "NAME"
            elif entity_type == "address":
                self.entity_map[key] = "ADDRESS"
            elif entity_type == "email":
                self.entity_map[key] = "EMAIL_ADDRESS"
            elif entity_type == "phone":
                self.entity_map[key] = "PHONE_NUMBER"
            elif entity_type == "ssn":
                self.entity_map[key] = "SSN"
            elif entity_type == "credit_card":
                self.entity_map[key] = "CREDIT_CARD_NUMBER"
            elif entity_type == "passport":
                self.entity_map[key] = "PASSPORT_NUMBER"
            elif entity_type == "eu_ssn":
                self.entity_map[key] = "EU_SSN"
            elif entity_type == "ip_address":
                self.entity_map[key] = "IP_ADDRESS"
            elif entity_type == "date":
                self.entity_map[key] = "DATE"
            else:
                self.entity_map[key] = f"ANONYMIZED_{entity_type.upper()}"
        return self.entity_map[key]
    
    def anonymize_credit_card(self, card: str) -> str:
        digits = re.sub(r'\D', '', card)
        if len(digits) >= 16:
            return "**** **** **** " + digits[-4:]
        elif len(digits) >= 13:
            return "**** ****** " + digits[-5:]
        elif len(digits) >= 4:
            return "*" * (len(digits) - 4) + digits[-4:]
        return "[ANONYMIZED_CARD]"
    
    def anonymize_name(self, name: str) -> str:
        return self.get_consistent_pseudonym("person", name)
    
    def anonymize_email(self, email: str) -> str:
        if '@' not in email:
            return email
        username, domain = email.split('@', 1)
        pseudonym = self.get_consistent_pseudonym("email", username)
        return f"{pseudonym}@{domain}"
    
    def anonymize_phone(self, phone: str) -> str:
        """Enhanced phone number anonymization"""
        # Remove all non-digit characters
        digits = re.sub(r'\D', '', phone)
        
        if len(digits) >= 10:  # Standard US phone number
            # Keep last 4 digits, mask the rest
            if len(digits) == 10:
                return f"***-***-{digits[-4:]}"
            elif len(digits) == 11 and digits[0] == '1':  # US with country code
                return f"***-***-{digits[-4:]}"
            else:
                return f"***-***-{digits[-4:]}"
        elif len(digits) >= 7:  # Partial phone number
            return f"***-***-{digits[-4:] if len(digits) >= 4 else digits}"
        else:
            # For very short numbers, just mask them
            return "*" * len(phone)
    
    def anonymize_ssn(self, ssn: str) -> str:
        """Anonymize SSN numbers"""
        if not ssn or not ssn.strip():
            return ssn
        return "[SSN]"
    
    def anonymize_address(self, address: str) -> str:
        return self.get_consistent_pseudonym("address", address)
    
    def anonymize_date(self, date_str: str) -> str:
        try:
            parsed_date = date_parser.parse(date_str, fuzzy=True)
            return "MM/DD/YYYY"
        except:
            return "DATE"
    
    def anonymize_ip_address(self, ip: str) -> str:
        return "IP_ADDRESS"
    
    def anonymize_id(self, id_str: str) -> str:
        """Anonymize various ID fields"""
        if not id_str or not id_str.strip():
            return id_str
        
        # Generate a consistent pseudonym for IDs
        return self.get_consistent_pseudonym("id", id_str)
    
    def anonymize_salary(self, salary: str) -> str:
        """Anonymize salary information"""
        if not salary or not salary.strip():
            return salary
        
        # Extract numbers from salary
        import re
        numbers = re.findall(r'\d+', salary)
        if numbers:
            # Generate a random salary in a reasonable range
            import random
            base_salary = random.randint(30000, 150000)
            return f"${base_salary:,}"
        
        return "[ANONYMIZED_SALARY]"
    
    def anonymize_commission(self, commission: str) -> str:
        """Anonymize commission information"""
        if not commission or not commission.strip():
            return commission
        
        # Extract numbers from commission
        import re
        numbers = re.findall(r'\d+', commission)
        if numbers:
            # Generate a random commission
            import random
            base_commission = random.randint(1000, 10000)
            return f"${base_commission:,}"
        
        return "[ANONYMIZED_COMMISSION]"
    
    def anonymize_postal_code(self, postal_code: str) -> str:
        """Anonymize postal codes"""
        if not postal_code or not postal_code.strip():
            return postal_code
        
        # Generate a fake postal code
        import random
        fake_postal = f"{random.randint(10000, 99999)}"
        return fake_postal
    
    def anonymize_passport(self, passport: str) -> str:
        """Anonymize passport numbers"""
        if not passport or not passport.strip():
            return passport
        return "PASSPORT_NUMBER"
    
    def anonymize_credit_card_exp(self, exp_date: str) -> str:
        """Anonymize credit card expiration dates"""
        if not exp_date or not exp_date.strip():
            return exp_date
        return "MM/YY"
    
    def anonymize_credit_card(self, card: str) -> str:
        """Anonymize credit card numbers with proper masking"""
        if not card or not card.strip():
            return card
        import re
        # Keep last 4 digits, mask the rest
        digits = re.sub(r'\D', '', card)
        if len(digits) >= 4:
            return "**** **** **** " + digits[-4:]
        return "[ANONYMIZED_CREDIT_CARD]"
    
    def anonymize_employee_id(self, employee_id: str) -> str:
        """Anonymize employee ID numbers"""
        if not employee_id or not employee_id.strip():
            return employee_id
        return "[EMPLOYEE_ID]"
    
    def anonymize_department(self, department: str) -> str:
        """Anonymize department names"""
        if not department or not department.strip():
            return department
        return "[DEPARTMENT]"
    
    def anonymize_routing_number(self, routing: str) -> str:
        """Anonymize routing numbers"""
        if not routing or not routing.strip():
            return routing
        return "[ROUTING_NUMBER]"
    
    def anonymize_bank_account(self, account: str) -> str:
        """Anonymize bank account numbers"""
        if not account or not account.strip():
            return account
        return "[BANK_ACCOUNT]"
    
    def anonymize_eu_ssn(self, ssn: str) -> str:
        """Anonymize European SSN numbers"""
        if not ssn or not ssn.strip():
            return ssn
        return "[EU_SSN]"

def detect_sensitive_patterns(content: str) -> List[Dict[str, Any]]:
    """Detect sensitive patterns in content with enhanced CV support"""
    sensitive_fields = []
    
    print(f"🔍 DEBUG: Analyzing content with {len(content)} characters")
    
    # Basic pattern detection
    for pattern_type, pattern in SENSITIVE_PATTERNS.items():
        matches = re.findall(pattern, content, re.IGNORECASE)
        if matches:
            clean_matches = []
            for match in matches:
                if isinstance(match, tuple):
                    # Handle tuple matches (shouldn't happen with our clean patterns)
                    match_str = ''.join(str(part) for part in match if part)
                else:
                    match_str = str(match)
                
                if match_str.strip():
                    clean_matches.append(match_str.strip())
            
            if clean_matches:
                sensitive_fields.append({
                    "field_type": pattern_type,
                    "count": len(clean_matches),
                    "examples": clean_matches[:10],  # Increased from 5 to 10
                    "risk_level": "high" if pattern_type in ["ssn", "credit_card", "cv_ssn", "cv_credit", "cv_bank"] else "medium"
                })
                print(f"🔍 DEBUG: Found {len(clean_matches)} matches for {pattern_type}")
    
    # Enhanced name detection - more precise for CVs
    all_names = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', content)
    person_names = []
    address_keywords = {'street', 'avenue', 'road', 'drive', 'lane', 'court', 'place', 'square', 'circle', 'terrace', 'main', 'oak', 'pine', 'elm', 'maple', 'north', 'south', 'east', 'west', 'central', 'park', 'view', 'hill', 'valley', 'bridge', 'station', 'airport', 'hospital', 'school', 'university', 'college', 'center', 'plaza', 'mall', 'store', 'shop', 'office', 'building', 'complex', 'tower', 'apartment', 'condo', 'house', 'home', 'residence', 'address', 'location', 'place', 'area', 'district', 'neighborhood', 'community', 'city', 'town', 'village', 'county', 'state', 'province', 'country', 'region', 'zone', 'section', 'block', 'lot', 'suite', 'floor', 'room', 'unit'}
    
    for name in all_names:
        name_parts = name.split()
        if len(name_parts) >= 2:
            # Check if any part looks like an address keyword
            is_address = any(part.lower() in address_keywords for part in name_parts)
            # Check if it looks like a person name (first letter capitalized, reasonable length)
            is_person = all(len(part) >= 2 and part[0].isupper() and part[1:].islower() for part in name_parts)
            
            if is_person and not is_address:
                person_names.append(name)
    
    if person_names:
        sensitive_fields.append({
            "field_type": "name",
            "count": len(person_names),
            "examples": person_names[:15],  # Increased from 10 to 15
            "risk_level": "high"
        })
        print(f"🔍 DEBUG: Found {len(person_names)} person names")
    
    # Add field-specific patterns for CSV fields
    csv_patterns = {
        'employee_id': r'\b\d{3,4}\b',  # Employee IDs like 198, 199, 200
        'job_id': r'\b[A-Z]{2,3}_[A-Z]+\b',  # Job IDs like SH_CLERK, AD_ASST
        'department_id': r'\b\d{1,3}\b',  # Department IDs like 50, 10, 20
        'manager_id': r'\b\d{1,3}\b',  # Manager IDs
        'salary': r'\b\d{4,5}\b',  # Salary amounts like 2600, 4400, 13000
        'commission': r'\b[-]\b',  # Commission fields
    }
    
    for field_type, pattern in csv_patterns.items():
        matches = re.findall(pattern, content, re.IGNORECASE)
        if matches:
            clean_matches = []
            for match in matches:
                if isinstance(match, tuple):
                    match_str = ''.join(str(part) for part in match if part)
                else:
                    match_str = str(match)
                
                if match_str.strip():
                    clean_matches.append(match_str.strip())
            
            if clean_matches:
                sensitive_fields.append({
                    "field_type": field_type,
                    "count": len(clean_matches),
                    "examples": clean_matches[:10],
                    "risk_level": "medium"
                })
    
    print(f"🔍 DEBUG: Total sensitive fields detected: {len(sensitive_fields)}")
    return sensitive_fields

def anonymize_csv_content(content: str, sensitive_fields: List[Dict], use_ai: bool = False) -> str:
    """GDPR-compliant CSV anonymization using Presidio"""
    try:
        # Get GDPR anonymizer
        gdpr_anonymizer = get_gdpr_anonymizer()
        
        csv_reader = csv.reader(io.StringIO(content))
        rows = list(csv_reader)
        
        if not rows:
            return content
        
        headers = rows[0]
        data_rows = rows[1:]
        
        # Convert to list of dictionaries for GDPR anonymization
        csv_data = []
        for row in data_rows:
            row_dict = {}
            for i, header in enumerate(headers):
                row_dict[header] = row[i] if i < len(row) else ""
            csv_data.append(row_dict)
        
        # Use GDPR-compliant anonymization
        anonymized_data = gdpr_anonymizer.anonymize_csv_data(csv_data)
        
        # Convert back to CSV format
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(headers)
        
        for row in anonymized_data:
            writer.writerow([row.get(header, "") for header in headers])
        
        return output.getvalue()
        
    except Exception as e:
        print(f"🔍 DEBUG: Error in GDPR CSV anonymization: {e}")
        # Fallback to basic anonymization
        return basic_anonymize_csv_content(content, sensitive_fields)

def basic_anonymize_csv_content(content: str, sensitive_fields: List[Dict]) -> str:
    """Fallback basic CSV anonymization"""
    try:
        engine = AnonymizationEngine()
        csv_reader = csv.reader(io.StringIO(content))
        rows = list(csv_reader)
        
        if not rows:
            return content
        
        headers = rows[0]
        data_rows = rows[1:]
        anonymized_rows = [headers]
        
        # Create comprehensive pattern matchers for each field type
        pattern_matchers = {}
        for field in sensitive_fields:
            field_type = field["field_type"]
            if field_type in SENSITIVE_PATTERNS:
                pattern_matchers[field_type] = re.compile(SENSITIVE_PATTERNS[field_type], re.IGNORECASE)
        
        # Add specific patterns for CSV fields
        csv_patterns = {
            'employee_id': re.compile(r'\b\d{3,4}\b', re.IGNORECASE),
            'job_id': re.compile(r'\b[A-Z]{2,3}_[A-Z]+\b', re.IGNORECASE),
            'department_id': re.compile(r'\b\d{1,3}\b', re.IGNORECASE),
            'manager_id': re.compile(r'\b\d{1,3}\b', re.IGNORECASE),
            'salary': re.compile(r'\b\d{4,5}\b', re.IGNORECASE),
        }
        pattern_matchers.update(csv_patterns)
        
        print(f"🔍 DEBUG: Processing {len(data_rows)} rows with {len(pattern_matchers)} pattern types")
        
        for row_index, row in enumerate(data_rows):
            anonymized_row = []
            
            for col_index, cell in enumerate(row):
                cell_str = str(cell).strip()
                anonymized_cell = cell_str
                
                # Check this cell against ALL patterns systematically
                for field_type, pattern_compiled in pattern_matchers.items():
                    if pattern_compiled.search(cell_str):
                        print(f"🔍 DEBUG: Row {row_index+1}, Col {col_index+1}: '{cell_str}' matches {field_type}")
                        
                        # This cell matches a sensitive pattern
                        if use_ai and ai_config["enabled"] and ai_config["api_key"]:
                            try:
                                # Use real AI-powered anonymization with CSV context
                                csv_context = engine.context_analyzer.analyze_csv_context(headers, row)
                                
                                # Create async event loop for AI call
                                import asyncio
                                try:
                                    loop = asyncio.get_event_loop()
                                except RuntimeError:
                                    loop = asyncio.new_event_loop()
                                    asyncio.set_event_loop(loop)
                                
                                # Call AI anonymization
                                anonymized_cell = loop.run_until_complete(
                                    engine.real_ai_anonymize(
                                        cell_str, 
                                        field_type, 
                                        str(csv_context), 
                                        headers, 
                                        col_index
                                    )
                                )
                                
                                print(f"🔍 DEBUG: AI generated: '{cell_str}' -> '{anonymized_cell}'")
                                
                            except Exception as ai_error:
                                print(f"🔍 DEBUG: AI failed, falling back to traditional: {ai_error}")
                                # Fallback to traditional method
                                anonymized_cell = _apply_traditional_anonymization(engine, cell_str, field_type)
                        elif use_ai:
                            # Use AI-powered anonymization with CSV context (legacy method)
                            csv_context = engine.context_analyzer.analyze_csv_context(headers, row)
                            anonymized_cell = engine.ai_anonymize_with_context(cell_str, field_type, csv_context)
                        else:
                            # Use traditional anonymization
                            anonymized_cell = _apply_traditional_anonymization(engine, cell_str, field_type)
                        
                        print(f"🔍 DEBUG: Anonymized '{cell_str}' -> '{anonymized_cell}'")
                        break  # Break after first match to avoid double-processing
                
                anonymized_row.append(anonymized_cell)
            
            anonymized_rows.append(anonymized_row)
        
        # Convert back to CSV string with proper formatting
        output = io.StringIO()
        csv_writer = csv.writer(output, lineterminator='\n')  # Ensure proper line endings
        csv_writer.writerows(anonymized_rows)
        csv_content = output.getvalue()
        
        # Ensure the CSV ends with a newline for proper formatting
        if not csv_content.endswith('\n'):
            csv_content += '\n'
        
        print(f"🔍 DEBUG: Successfully anonymized CSV with {len(anonymized_rows)} rows")
        return csv_content
        
    except Exception as e:
        print(f"Error anonymizing CSV: {e}")
        import traceback
        traceback.print_exc()
        return content

def anonymize_text_content(content: str, sensitive_fields: List[Dict], use_ai: bool = False) -> str:
    """GDPR-compliant text content anonymization using Presidio with traditional fallback"""
    try:
        # Get GDPR anonymizer
        gdpr_anonymizer = get_gdpr_anonymizer()
        
        # Use GDPR-compliant anonymization
        result = gdpr_anonymizer.anonymize_text(content)
        
        if result.get("error"):
            print(f"🔍 DEBUG: GDPR text anonymization failed: {result['error']}")
            # Fallback to basic anonymization
            return basic_anonymize_text_content(content, sensitive_fields, use_ai=use_ai)
        
        print(f"🔍 DEBUG: GDPR text anonymization successful - found {result['total_entities']} entities")
        
        # Apply traditional anonymization as additional layer for patterns not caught by GDPR
        gdpr_anonymized = result["anonymized_text"]
        traditional_anonymized = basic_anonymize_text_content(gdpr_anonymized, sensitive_fields, use_ai=use_ai)
        
        return traditional_anonymized
        
    except Exception as e:
        print(f"🔍 DEBUG: Error in GDPR text anonymization: {e}")
        # Fallback to basic anonymization
        return basic_anonymize_text_content(content, sensitive_fields, use_ai=use_ai)

def basic_anonymize_text_content(content: str, sensitive_fields: List[Dict], use_ai: bool = False) -> str:
    """Fallback basic text content anonymization"""
    try:
        anonymized_content = content
        engine = AnonymizationEngine()
        
        # First, let's use regex patterns to find and replace sensitive data
        import re
        
        # Define patterns for different types of sensitive data
        patterns = {
            'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'phone': r'(\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}',
            'ssn': r'\b\d{3}-?\d{2}-?\d{4}\b',
            'credit_card': r'\b(?:\d{4}[-\s]?){3}\d{4}\b',
            'credit_card_exp': r'\b(?:0[1-9]|1[0-2])[/-]\d{2}\b',
            'passport': r'\b[A-Z]{2,3}\d{6,9}\b',
            'ip_address': r'\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b',
            'date': r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
            'address': r'\b\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Lane|Ln|Drive|Dr|Boulevard|Blvd|Way|Place|Pl|Court|Ct|Circle|Cir)\b',
            'postal_code': r'\b[A-Z]{1,2}\d[A-Z\d]?\s?\d[A-Z]{2}\b'
        }
        
        # Process each field type
        for field in sensitive_fields:
            field_type = field["field_type"]
            
            if field_type in ["email", "Email"]:
                # Find and replace email addresses
                emails = re.findall(patterns['email'], anonymized_content)
                for email in emails:
                    if use_ai:
                        replacement = engine.ai_anonymize_with_context(email, field_type, "")
                    else:
                        replacement = engine.anonymize_email(email)
                    anonymized_content = anonymized_content.replace(email, replacement)
                    
            elif field_type in ["phone", "Phone"]:
                # Find and replace phone numbers
                phones = re.findall(patterns['phone'], anonymized_content)
                for phone in phones:
                    if use_ai:
                        replacement = engine.ai_anonymize_with_context(phone, field_type, "")
                    else:
                        replacement = engine.anonymize_phone(phone)
                    anonymized_content = anonymized_content.replace(phone, replacement)
                    
            elif field_type in ["name", "person", "Name", "Person"]:
                # For names, we need to be more careful - look for common name patterns
                # This is a simplified approach - in practice, you'd want more sophisticated name detection
                name_pattern = r'\b[A-Z][a-z]+ [A-Z][a-z]+\b'
                names = re.findall(name_pattern, anonymized_content)
                for name in names:
                    if use_ai:
                        replacement = engine.ai_anonymize_with_context(name, field_type, "")
                    else:
                        replacement = engine.anonymize_name(name)
                    anonymized_content = anonymized_content.replace(name, replacement)
                    
            elif field_type in ["credit_card", "CreditCard"]:
                # Find and replace credit card numbers
                cards = re.findall(patterns['credit_card'], anonymized_content)
                for card in cards:
                    if use_ai:
                        replacement = engine.ai_anonymize_with_context(card, field_type, "")
                    else:
                        replacement = engine.anonymize_credit_card(card)
                    anonymized_content = anonymized_content.replace(card, replacement)
                    
            elif field_type in ["ssn", "SSN"]:
                # Find and replace SSNs
                ssns = re.findall(patterns['ssn'], anonymized_content)
                for ssn in ssns:
                    if use_ai:
                        replacement = engine.ai_anonymize_with_context(ssn, field_type, "")
                    else:
                        replacement = engine.anonymize_ssn(ssn)
                    anonymized_content = anonymized_content.replace(ssn, replacement)
                    
            elif field_type in ["eu_ssn", "EU_SSN"]:
                # Find and replace European SSNs
                eu_ssns = re.findall(patterns['eu_ssn'], anonymized_content)
                for eu_ssn in eu_ssns:
                    if use_ai:
                        replacement = engine.ai_anonymize_with_context(eu_ssn, field_type, "")
                    else:
                        replacement = "[EU_SSN]"
                    anonymized_content = anonymized_content.replace(eu_ssn, replacement)
                    
            elif field_type in ["employee_id", "Employee_ID"]:
                # Find and replace Employee IDs
                employee_ids = re.findall(patterns['employee_id'], anonymized_content)
                for emp_id in employee_ids:
                    if use_ai:
                        replacement = engine.ai_anonymize_with_context(emp_id, field_type, "")
                    else:
                        replacement = "[EMPLOYEE_ID]"
                    anonymized_content = anonymized_content.replace(emp_id, replacement)
                    
            elif field_type in ["salary", "Salary"]:
                # Find and replace Salary amounts
                salaries = re.findall(patterns['salary'], anonymized_content)
                for salary in salaries:
                    if use_ai:
                        replacement = engine.ai_anonymize_with_context(salary, field_type, "")
                    else:
                        replacement = "[SALARY]"
                    anonymized_content = anonymized_content.replace(salary, replacement)
                    
            elif field_type in ["department", "Department"]:
                # Find and replace Department names
                departments = re.findall(patterns['department'], anonymized_content)
                for dept in departments:
                    if use_ai:
                        replacement = engine.ai_anonymize_with_context(dept, field_type, "")
                    else:
                        replacement = "[DEPARTMENT]"
                    anonymized_content = anonymized_content.replace(dept, replacement)
                    
            elif field_type in ["routing_number", "Routing_Number"]:
                # Find and replace Routing Numbers
                routing_numbers = re.findall(patterns['routing_number'], anonymized_content)
                for routing in routing_numbers:
                    if use_ai:
                        replacement = engine.ai_anonymize_with_context(routing, field_type, "")
                    else:
                        replacement = "[ROUTING_NUMBER]"
                    anonymized_content = anonymized_content.replace(routing, replacement)
                    
            elif field_type in ["ip_address", "IPAddress"]:
                # Find and replace IP addresses
                ips = re.findall(patterns['ip_address'], anonymized_content)
                for ip in ips:
                    if use_ai:
                        replacement = engine.ai_anonymize_with_context(ip, field_type, "")
                    else:
                        replacement = engine.anonymize_ip_address(ip)
                    anonymized_content = anonymized_content.replace(ip, replacement)
                    
            elif field_type in ["date", "Date"]:
                # Find and replace dates
                dates = re.findall(patterns['date'], anonymized_content)
                for date in dates:
                    if use_ai:
                        replacement = engine.ai_anonymize_with_context(date, field_type, "")
                    else:
                        replacement = engine.anonymize_date(date)
                    anonymized_content = anonymized_content.replace(date, replacement)
                    
            elif field_type in ["passport", "Passport"]:
                # Find and replace passport numbers
                passports = re.findall(patterns['passport'], anonymized_content)
                for passport in passports:
                    if use_ai:
                        replacement = engine.ai_anonymize_with_context(passport, field_type, "")
                    else:
                        replacement = engine.anonymize_passport(passport)
                    anonymized_content = anonymized_content.replace(passport, replacement)
                    
            elif field_type in ["address", "Address"]:
                # Find and replace addresses
                addresses = re.findall(patterns['address'], anonymized_content)
                for address in addresses:
                    if use_ai:
                        replacement = engine.ai_anonymize_with_context(address, field_type, "")
                    else:
                        replacement = engine.anonymize_address(address)
                    anonymized_content = anonymized_content.replace(address, replacement)
                    
            elif field_type in ["postal_code", "PostalCode"]:
                # Find and replace postal codes
                postal_codes = re.findall(patterns['postal_code'], anonymized_content)
                for postal_code in postal_codes:
                    if use_ai:
                        replacement = engine.ai_anonymize_with_context(postal_code, field_type, "")
                    else:
                        replacement = engine.anonymize_postal_code(postal_code)
                    anonymized_content = anonymized_content.replace(postal_code, replacement)
                    
            elif field_type in ["credit_card_exp", "CreditCardExp"]:
                # Find and replace credit card expiration dates
                exp_dates = re.findall(patterns['credit_card_exp'], anonymized_content)
                for exp_date in exp_dates:
                    if use_ai:
                        replacement = engine.ai_anonymize_with_context(exp_date, field_type, "")
                    else:
                        replacement = engine.anonymize_credit_card_exp(exp_date)
                    anonymized_content = anonymized_content.replace(exp_date, replacement)
                    
            else:
                # For other field types, use the original approach but clean up examples
                examples = field["examples"]
                for example in examples:
                    # Clean up the example by removing newlines and extra whitespace
                    clean_example = ' '.join(example.split())
                    if clean_example and len(clean_example) > 2:  # Only process meaningful examples
                        if clean_example in anonymized_content:
                            if use_ai:
                                replacement = engine.ai_anonymize_with_context(clean_example, field_type, "")
                            else:
                                replacement = f"[ANONYMIZED_{field_type.upper()}]"
                            anonymized_content = anonymized_content.replace(clean_example, replacement)
        
        return anonymized_content
        
    except Exception as e:
        print(f"Error anonymizing text: {e}")
        return content

def _apply_traditional_anonymization(engine: AnonymizationEngine, cell_str: str, field_type: str) -> str:
    """Apply traditional anonymization based on field type"""
    try:
        if field_type in ["email", "Email"]:
            return engine.anonymize_email(cell_str)
        elif field_type in ["phone", "Phone"]:
            return engine.anonymize_phone(cell_str)
        elif field_type in ["ssn", "SSN"]:
            return engine.anonymize_ssn(cell_str)
        elif field_type in ["name", "person", "Name", "Person"]:
            return engine.anonymize_name(cell_str)
        elif field_type in ["address", "Address"]:
            return engine.anonymize_address(cell_str)
        elif field_type in ["credit_card", "CreditCard"]:
            return engine.anonymize_credit_card(cell_str)
        elif field_type in ["date", "Date"]:
            return engine.anonymize_date(cell_str)
        elif field_type in ["expiry_date", "ExpiryDate"]:
            return engine.anonymize_expiry_date(cell_str)
        elif field_type in ["ip_address", "IPAddress"]:
            return engine.anonymize_ip_address(cell_str)
        elif field_type in ["mac_address", "MacAddress"]:
            return engine.anonymize_mac_address(cell_str)
        elif field_type in ["vehicle_plate", "VehiclePlate"]:
            return engine.anonymize_vehicle_plate(cell_str)
        elif field_type in ["bank_account", "BankAccount"]:
            return engine.anonymize_bank_account(cell_str)
        elif field_type in ["postal_code", "PostalCode"]:
            return engine.anonymize_postal_code(cell_str)
        elif field_type in ["salary", "Salary"]:
            return engine.anonymize_salary(cell_str)
        elif field_type in ["employee_id", "EmployeeID"]:
            return engine.anonymize_id(cell_str)
        elif field_type in ["job_id", "JobID"]:
            return engine.anonymize_id(cell_str)
        elif field_type in ["department_id", "DepartmentID"]:
            return engine.anonymize_id(cell_str)
        elif field_type in ["manager_id", "ManagerID"]:
            return engine.anonymize_id(cell_str)
        elif field_type in ["commission", "Commission"]:
            return engine.anonymize_commission(cell_str)
        else:
            return f"[ANONYMIZED_{field_type.upper()}]"
    except Exception as e:
        print(f"🔍 DEBUG: Error in traditional anonymization for {field_type}: {e}")
        return f"[ANONYMIZED_{field_type.upper()}]"

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Extract text from various file types with enhanced Word document support"""
    file_extension = os.path.splitext(filename)[1].lower()
    
    try:
        if file_extension == '.csv':
            return file_content.decode('utf-8', errors='ignore')
        elif file_extension in ['.txt', '.json']:
            return file_content.decode('utf-8', errors='ignore')
        elif file_extension == '.docx':
            # Enhanced Word document text extraction
            import docx
            doc = docx.Document(io.BytesIO(file_content))
            text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(' | '.join(row_text))
            
            # Extract text from headers and footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text.append(paragraph.text.strip())
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text.append(paragraph.text.strip())
            
            result = '\n'.join(text)
            print(f"🔍 DEBUG: Extracted {len(result)} characters from Word document")
            return result
            
        elif file_extension == '.pdf':
            # Handle PDF files
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = []
            for page in pdf_reader.pages:
                text.append(page.extract_text())
            return '\n'.join(text)
        elif file_extension in ['.xlsx', '.xls']:
            # Handle Excel files
            import pandas as pd
            df = pd.read_excel(io.BytesIO(file_content))
            return df.to_csv(index=False)
        else:
            # Fallback: try to decode as text
            return file_content.decode('utf-8', errors='ignore')
    except Exception as e:
        print(f"🔍 DEBUG: Error extracting text from {filename}: {e}")
        # Fallback: return a placeholder
        return f"Error extracting text from {filename}: {str(e)}"

# API Endpoints
@app.get("/api/v1/health")
async def health_check():
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "service": "GDPR-Compliant Data Anonymization Service",
        "version": "2.0.0",
        "features": {
            "gdpr_compliance": True,
            "rate_limiting": True,
            "file_validation": True,
            "progress_tracking": True
        }
    }

@app.get("/api/v1/status/{file_id}")
async def get_processing_status(file_id: str):
    """Get processing status for a file"""
    if file_id not in processing_status:
        raise HTTPException(status_code=404, detail="File not found")
    
    return processing_status[file_id]

@app.get("/api/v1/stats")
async def get_service_stats():
    """Get service statistics"""
    return {
        "total_files_processed": len(uploaded_files),
        "active_processing": len([s for s in processing_status.values() if s["status"] == "processing"]),
        "rate_limit_requests": sum(len(requests) for requests in rate_limit_storage.values()),
        "uptime": datetime.now().isoformat()
    }

@app.post("/api/v1/batch/upload")
async def batch_upload_files(request: Request, files: List[UploadFile] = File(...)):
    """Upload multiple files at once"""
    try:
        # Rate limiting
        check_rate_limit(request)
        
        if len(files) > 10:  # Limit batch size
            raise HTTPException(
                status_code=400, 
                detail="Too many files. Maximum 10 files per batch."
            )
        
        batch_id = str(uuid.uuid4())
        batch_results = []
        
        for file in files:
            try:
                # File validation
                validate_file_extension(file.filename)
                validate_file_size(file)
                
                file_id = str(uuid.uuid4())
                content = await file.read()
                
                # Additional size check
                if len(content) > MAX_FILE_SIZE:
                    batch_results.append({
                        "filename": file.filename,
                        "status": "error",
                        "error": f"File too large. Maximum size allowed is {MAX_FILE_SIZE // (1024*1024)}MB"
                    })
                    continue
                
                text_content = extract_text_from_file(content, file.filename)
                file_extension = os.path.splitext(file.filename)[1].lower()
                
                if file_extension == '.csv':
                    file_type = "csv"
                elif file_extension in ['.xlsx', '.xls']:
                    file_type = "excel"
                elif file_extension == '.docx':
                    file_type = "docx"
                elif file_extension == '.pdf':
                    file_type = "pdf"
                elif file_extension == '.json':
                    file_type = "json"
                else:
                    file_type = "text"
                
                uploaded_files[file_id] = {
                    "filename": file.filename,
                    "content": content,
                    "text_content": text_content,
                    "file_type": file_type,
                    "uploaded_at": datetime.now().isoformat(),
                    "batch_id": batch_id
                }
                
                batch_results.append({
                    "filename": file.filename,
                    "file_id": file_id,
                    "status": "success",
                    "file_type": file_type
                })
                
            except Exception as e:
                batch_results.append({
                    "filename": file.filename,
                    "status": "error",
                    "error": str(e)
                })
        
        return {
            "batch_id": batch_id,
            "total_files": len(files),
            "successful_uploads": len([r for r in batch_results if r["status"] == "success"]),
            "failed_uploads": len([r for r in batch_results if r["status"] == "error"]),
            "results": batch_results
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Batch upload failed: {str(e)}")

@app.get("/api/v1/ai-config")
async def get_ai_config():
    """Get current AI configuration"""
    return {
        "enabled": ai_config["enabled"],
        "provider": ai_config["provider"],
        "model": ai_config["model"],
        "has_api_key": ai_config["api_key"] is not None,
        "base_url": ai_config["base_url"],
        "temperature": ai_config["temperature"],
        "max_tokens": ai_config["max_tokens"],
        "fallback_to_traditional": ai_config["fallback_to_traditional"]
    }

@app.post("/api/v1/ai-config")
async def update_ai_config(request: Request):
    """Update AI configuration"""
    try:
        data = await request.json()
        
        if "api_key" in data:
            ai_config["api_key"] = data["api_key"]
        if "model" in data:
            ai_config["model"] = data["model"]
        if "temperature" in data:
            ai_config["temperature"] = data["temperature"]
        if "max_tokens" in data:
            ai_config["max_tokens"] = data["max_tokens"]
        if "enabled" in data:
            ai_config["enabled"] = data["enabled"]
        if "provider" in data:
            ai_config["provider"] = data["provider"]
        if "base_url" in data:
            ai_config["base_url"] = data["base_url"]
        if "fallback_to_traditional" in data:
            ai_config["fallback_to_traditional"] = data["fallback_to_traditional"]
        
        return {
            "status": "updated",
            "config": ai_config
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to update AI config: {str(e)}")

@app.post("/api/v1/upload")
async def upload_file(request: Request, file: UploadFile = File(...)):
    try:
        # Rate limiting
        check_rate_limit(request)
        
        # File validation
        validate_file_extension(file.filename)
        validate_file_size(file)
        
        file_id = str(uuid.uuid4())
        content = await file.read()
        
        # Additional size check after reading
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size allowed is {MAX_FILE_SIZE // (1024*1024)}MB"
            )
        
        text_content = extract_text_from_file(content, file.filename)
        
        file_extension = os.path.splitext(file.filename)[1].lower()
        print(f"🔍 DEBUG: File extension detected: {file_extension}")
        if file_extension == '.csv':
            file_type = "csv"
            # For CSV files, preserve the original content as CSV
            content_to_store = content.decode('utf-8', errors='ignore')
        elif file_extension in ['.xlsx', '.xls']:
            file_type = "excel"
            # For Excel files, we need to preserve the binary content for proper output
            content_to_store = content  # Keep as bytes for Excel
        elif file_extension == '.pdf':
            file_type = "pdf"
            # For PDF files, we need to preserve the binary content for proper output
            content_to_store = content  # Keep as bytes for PDF
        elif file_extension == '.docx':
            file_type = "docx"
            # For Word files, we need to preserve the binary content for proper output
            content_to_store = content  # Keep as bytes for Word
        elif file_extension == '.json':
            file_type = "json"
            content_to_store = text_content
        elif file_extension == '.txt':
            file_type = "text"
            content_to_store = text_content
        else:
            file_type = "unknown"
            content_to_store = content  # Keep as bytes for unknown types
        
        uploaded_files[file_id] = {
            "filename": file.filename,
            "content": content_to_store,
            "original_content": content,
            "file_type": file_type,
            "size": len(content),
            "uploaded_at": datetime.now().isoformat()
        }
        
        return {
            "file_id": file_id,
            "filename": file.filename,
            "status": "uploaded",
            "file_type": file_type,
            "size": len(content)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@app.post("/api/v1/analyze")
async def analyze_file(request: Request):
    try:
        data = await request.json()
        file_id = data.get("file_id")
        
        if file_id not in uploaded_files:
            raise HTTPException(status_code=404, detail="File not found")
        
        file_data = uploaded_files[file_id]
        content = file_data["content"]
        file_type = file_data["file_type"]
        
        # Ensure content is text for analysis
        if isinstance(content, bytes):
            print(f"🔍 DEBUG: Content is bytes, file_type: {file_type}")
            if file_type == "docx":
                # Extract text from Word document
                import docx
                doc = docx.Document(io.BytesIO(content))
                text_content = []
                for paragraph in doc.paragraphs:
                    text_content.append(paragraph.text)
                content = '\n'.join(text_content)
                print(f"🔍 DEBUG: Extracted text from Word: {content[:100]}...")
            elif file_type == "pdf":
                # Extract text from PDF
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                text_content = []
                for page in pdf_reader.pages:
                    text_content.append(page.extract_text())
                content = '\n'.join(text_content)
            elif file_type == "excel":
                # Extract text from Excel
                import pandas as pd
                df = pd.read_excel(io.BytesIO(content))
                content = df.to_csv(index=False)
            else:
                # Fallback: decode as text
                content = content.decode('utf-8', errors='ignore')
        
        try:
            sensitive_fields = detect_sensitive_patterns(content)
        except Exception as e:
            print(f"🔍 DEBUG: Error in detect_sensitive_patterns: {e}")
            raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")
        
        total_sensitive = sum(field["count"] for field in sensitive_fields)
        if total_sensitive > 10:
            risk_level = "HIGH"
        elif total_sensitive > 5:
            risk_level = "MEDIUM"
        else:
            risk_level = "LOW"
        
        uploaded_files[file_id]["analysis_results"] = {
            "sensitive_fields": sensitive_fields,
            "total_count": total_sensitive,
            "risk_level": risk_level,
            "analyzed_at": datetime.now().isoformat()
        }
        
        return {
            "file_id": file_id,
            "status": "analyzed",
            "file_type": file_type,
            "sensitive_fields": sensitive_fields,
            "total_count": total_sensitive,
            "risk_level": risk_level
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

def recreate_excel_file(content: bytes, filename: str, sensitive_fields: list, use_ai: bool = False) -> bytes:
    """Recreate Excel file with anonymized content"""
    try:
        import pandas as pd
        import io
        
        # Read the original Excel file
        df = pd.read_excel(io.BytesIO(content))
        
        # Anonymize the DataFrame
        anonymized_df = anonymize_dataframe(df, sensitive_fields, use_ai=use_ai)
        
        # Convert back to Excel format
        output = io.BytesIO()
        anonymized_df.to_excel(output, index=False, engine='openpyxl')
        return output.getvalue()
        
    except Exception as e:
        print(f"Error recreating Excel file: {e}")
        # Fallback to CSV
        text_content = extract_text_from_file(content, filename)
        anonymized_text = anonymize_text_content(text_content, sensitive_fields, use_ai=use_ai)
        return anonymized_text.encode('utf-8')

def recreate_docx_file(content: bytes, filename: str, sensitive_fields: list, use_ai: bool = False) -> bytes:
    """Recreate Word document with anonymized content"""
    try:
        from docx import Document
        import io
        
        # Read the original Word document
        doc = Document(io.BytesIO(content))
        
        # Anonymize all paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                anonymized_text = anonymize_text_content(paragraph.text, sensitive_fields, use_ai=use_ai)
                paragraph.text = anonymized_text
        
        # Anonymize all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        anonymized_text = anonymize_text_content(cell.text, sensitive_fields, use_ai=use_ai)
                        cell.text = anonymized_text
        
        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
        
    except Exception as e:
        print(f"Error recreating Word file: {e}")
        # Fallback to text
        text_content = extract_text_from_file(content, filename)
        anonymized_text = anonymize_text_content(text_content, sensitive_fields, use_ai=use_ai)
        return anonymized_text.encode('utf-8')

def recreate_pdf_file(content: bytes, filename: str, sensitive_fields: list, use_ai: bool = False) -> bytes:
    """Recreate PDF file with anonymized content"""
    try:
        import PyPDF2
        import io
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        # Extract text from original PDF
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text_content = ""
        for page in pdf_reader.pages:
            text_content += page.extract_text() + "\n"
        
        # Anonymize the text
        anonymized_text = anonymize_text_content(text_content, sensitive_fields, use_ai=use_ai)
        
        # Create new PDF with anonymized content
        output = io.BytesIO()
        c = canvas.Canvas(output, pagesize=letter)
        
        # Split text into lines and add to PDF
        lines = anonymized_text.split('\n')
        y_position = 750
        for line in lines[:50]:  # Limit to first 50 lines to avoid overflow
            if y_position < 50:
                c.showPage()
                y_position = 750
            c.drawString(50, y_position, line[:80])  # Limit line length
            y_position -= 15
        
        c.save()
        return output.getvalue()
        
    except Exception as e:
        print(f"Error recreating PDF file: {e}")
        # Fallback to text
        text_content = extract_text_from_file(content, filename)
        anonymized_text = anonymize_text_content(text_content, sensitive_fields, use_ai=use_ai)
        return anonymized_text.encode('utf-8')

def anonymize_dataframe(df, sensitive_fields: list, use_ai: bool = False):
    """Anonymize a pandas DataFrame"""
    anonymized_df = df.copy()
    
    for field_info in sensitive_fields:
        field_type = field_info["field_type"]
        
        # Find columns that might contain this type of data
        for col in anonymized_df.columns:
            if anonymized_df[col].dtype == 'object':  # Text columns
                anonymized_df[col] = anonymized_df[col].astype(str).apply(
                    lambda x: anonymize_text_content(x, [field_info], use_ai=use_ai) if pd.notna(x) and x.strip() else x
                )
    
    return anonymized_df

@app.post("/api/v1/anonymize")
async def anonymize_file(request: Request):
    try:
        # Rate limiting
        check_rate_limit(request)
        
        data = await request.json()
        file_id = data.get("file_id")
        use_ai = data.get("use_ai", True)  # Default to AI-powered
        
        # Check if AI is enabled and configured
        if use_ai and not ai_config["enabled"]:
            use_ai = False
        if use_ai and not ai_config["api_key"]:
            use_ai = False
        
        if file_id not in uploaded_files:
            raise HTTPException(status_code=404, detail="File not found")
        
        # Initialize progress tracking
        processing_status[file_id] = {
            "status": "processing",
            "progress": 0,
            "message": "Starting anonymization...",
            "started_at": datetime.now().isoformat()
        }
        
        file_data = uploaded_files[file_id]
        content = file_data["content"]
        file_type = file_data["file_type"]
        
        if "analysis_results" not in file_data:
            raise HTTPException(status_code=400, detail="File must be analyzed before anonymization")
        
        sensitive_fields = file_data["analysis_results"]["sensitive_fields"]
        
        # Update progress
        processing_status[file_id].update({
            "progress": 25,
            "message": f"Anonymizing {file_type.upper()} content..."
        })
        
        if file_type == "csv":
            anonymized_content = anonymize_csv_content(content, sensitive_fields, use_ai=use_ai)
        elif file_type == "excel":
            # For Excel files, recreate the Excel format with anonymized data
            anonymized_content = recreate_excel_file(content, file_data["filename"], sensitive_fields, use_ai=use_ai)
        elif file_type == "docx":
            # For Word files, recreate the Word format with anonymized data
            anonymized_content = recreate_docx_file(content, file_data["filename"], sensitive_fields, use_ai=use_ai)
        elif file_type == "pdf":
            # For PDF files, recreate the PDF format with anonymized data
            anonymized_content = recreate_pdf_file(content, file_data["filename"], sensitive_fields, use_ai=use_ai)
        else:
            anonymized_content = anonymize_text_content(content, sensitive_fields, use_ai=use_ai)
        
        # Update progress
        processing_status[file_id].update({
            "progress": 75,
            "message": "Preparing anonymized file..."
        })
        
        anonymized_file_id = str(uuid.uuid4())
        
        # Ensure the anonymized filename has the correct extension
        original_filename = file_data['filename']
        if file_type == "csv" and not original_filename.endswith('.csv'):
            anonymized_filename = f"anonymized_{os.path.splitext(original_filename)[0]}.csv"
        elif file_type == "excel" and not original_filename.endswith('.xlsx'):
            anonymized_filename = f"anonymized_{os.path.splitext(original_filename)[0]}.xlsx"
        elif file_type == "text" and not original_filename.endswith('.txt'):
            anonymized_filename = f"anonymized_{os.path.splitext(original_filename)[0]}.txt"
        elif file_type == "json" and not original_filename.endswith('.json'):
            anonymized_filename = f"anonymized_{os.path.splitext(original_filename)[0]}.json"
        else:
            anonymized_filename = f"anonymized_{original_filename}"
        
        techniques = ["ai_powered_anonymization"] if use_ai else ["traditional_anonymization"]
        
        uploaded_files[anonymized_file_id] = {
            "filename": anonymized_filename,
            "content": anonymized_content,
            "file_type": file_type,
            "anonymized_at": datetime.now().isoformat(),
            "original_file_id": file_id,
            "techniques_applied": techniques
        }
        
        # Complete progress tracking
        processing_status[file_id].update({
            "status": "completed",
            "progress": 100,
            "message": "Anonymization completed successfully!",
            "completed_at": datetime.now().isoformat(),
            "anonymized_file_id": anonymized_file_id
        })
        
        return {
            "status": "completed",
            "anonymized_file_id": anonymized_file_id,
            "original_file_id": file_id,
            "techniques_applied": techniques,
            "anonymization_method": "ai_powered" if use_ai else "traditional"
        }
        
    except Exception as e:
        # Update progress with error
        if file_id in processing_status:
            processing_status[file_id].update({
                "status": "error",
                "progress": 0,
                "message": f"Anonymization failed: {str(e)}",
                "error": str(e),
                "failed_at": datetime.now().isoformat()
            })
        raise HTTPException(status_code=500, detail=f"Anonymization failed: {str(e)}")

@app.get("/api/v1/download/{file_id}")
async def download_file(file_id: str):
    try:
        if file_id not in uploaded_files:
            raise HTTPException(status_code=404, detail="File not found")
        
        file_data = uploaded_files[file_id]
        content = file_data["content"]
        filename = file_data["filename"]
        file_type = file_data["file_type"]
        
        print(f"🔍 DEBUG: Downloading file {file_id}")
        print(f"🔍 DEBUG: file_type = {file_type}")
        print(f"🔍 DEBUG: filename = {filename}")
        
        # Force the correct file extension based on file_type
        if file_type == "csv":
            # Ensure filename ends with .csv
            if not filename.endswith('.csv'):
                filename = filename.replace('.txt', '.csv').replace('.xlsx', '.csv').replace('.xls', '.csv')
                if not filename.endswith('.csv'):
                    filename = f"{filename}.csv"
            suffix = ".csv"
            media_type = "text/csv"
        elif file_type == "excel":
            if not filename.endswith('.xlsx'):
                filename = filename.replace('.csv', '.xlsx').replace('.txt', '.xlsx')
                if not filename.endswith('.xlsx'):
                    filename = f"{filename}.xlsx"
            suffix = ".xlsx"
            media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        elif file_type == "pdf":
            if not filename.endswith('.pdf'):
                filename = filename.replace('.txt', '.pdf').replace('.csv', '.pdf')
                if not filename.endswith('.pdf'):
                    filename = f"{filename}.pdf"
            suffix = ".pdf"
            media_type = "application/pdf"
        elif file_type == "docx":
            if not filename.endswith('.docx'):
                filename = filename.replace('.txt', '.docx').replace('.csv', '.docx')
                if not filename.endswith('.docx'):
                    filename = f"{filename}.docx"
            suffix = ".docx"
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif file_type == "text":
            if not filename.endswith('.txt'):
                filename = filename.replace('.csv', '.txt').replace('.xlsx', '.txt')
                if not filename.endswith('.txt'):
                    filename = f"{filename}.txt"
            suffix = ".txt"
            media_type = "text/plain"
        elif file_type == "json":
            if not filename.endswith('.json'):
                filename = filename.replace('.csv', '.json').replace('.txt', '.json')
                if not filename.endswith('.json'):
                    filename = f"{filename}.json"
            suffix = ".json"
            media_type = "application/json"
        else:
            # For unknown types, preserve original extension
            suffix = os.path.splitext(filename)[1]
            if not suffix:
                suffix = ".txt"
            media_type = "application/octet-stream"
        
        print(f"🔍 DEBUG: Final filename = {filename}")
        print(f"🔍 DEBUG: Using suffix = {suffix}")
        print(f"🔍 DEBUG: Using media_type = {media_type}")
        
        # Create temporary file with the correct extension
        if file_type in ["excel", "pdf", "docx"]:
            # For binary files, write as bytes
            temp_file = tempfile.NamedTemporaryFile(mode='wb', delete=False, suffix=suffix)
            if isinstance(content, str):
                temp_file.write(content.encode('utf-8'))
            else:
                temp_file.write(content)
            temp_file.close()
        else:
            # For text files, write as text
            temp_file = tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=suffix, encoding='utf-8')
            if isinstance(content, bytes):
                temp_file.write(content.decode('utf-8', errors='ignore'))
            else:
                temp_file.write(content)
            temp_file.close()
        
        print(f"🔍 DEBUG: Created temp file: {temp_file.name}")
        
        # Clean up temporary file after response
        import atexit
        def cleanup_temp_file():
            try:
                if os.path.exists(temp_file.name):
                    os.unlink(temp_file.name)
                    print(f"🔍 DEBUG: Cleaned up temp file: {temp_file.name}")
            except Exception as e:
                print(f"🔍 DEBUG: Cleanup error: {e}")
        
        atexit.register(cleanup_temp_file)
        
        response = FileResponse(
            temp_file.name,
            media_type=media_type,
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
        
        print(f"🔍 DEBUG: Returning response with filename: {filename}")
        return response
        
    except Exception as e:
        print(f"❌ Download error for file {file_id}: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
